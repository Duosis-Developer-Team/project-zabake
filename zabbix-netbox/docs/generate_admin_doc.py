import os
from pathlib import Path
from datetime import datetime

from docx import Document
from docx.shared import Inches
from fpdf import FPDF
import matplotlib.pyplot as plt


BASE_DIR = Path(__file__).resolve().parent.parent
DOCS_DIR = BASE_DIR / "zabbix-netbox" / "docs" if (BASE_DIR / "zabbix-netbox").exists() else Path(__file__).resolve().parent


def _ensure_output_dir() -> Path:
    output_dir = DOCS_DIR
    output_dir.mkdir(parents=True, exist_ok=True)
    return output_dir


def _generate_diagram(output_path: Path, title: str, boxes: list[str], arrows: list[tuple[int, int]] | None = None) -> None:
    fig, ax = plt.subplots(figsize=(8, 4))
    ax.set_axis_off()

    n = len(boxes)
    x_positions = [i / max(n - 1, 1) for i in range(n)]
    y = 0.5

    for idx, (label, x) in enumerate(zip(boxes, x_positions)):
        ax.text(
            x,
            y,
            label,
            ha="center",
            va="center",
            fontsize=10,
            bbox=dict(boxstyle="round,pad=0.4", edgecolor="black", facecolor="#e0f2f1"),
        )

    if arrows:
        for start, end in arrows:
            if 0 <= start < n and 0 <= end < n:
                x_start = x_positions[start]
                x_end = x_positions[end]
                ax.annotate(
                    "",
                    xy=(x_end - 0.03, y),
                    xytext=(x_start + 0.03, y),
                    arrowprops=dict(arrowstyle="->", lw=1.5),
                )

    ax.set_title(title, fontsize=12, pad=20)
    fig.tight_layout()
    fig.savefig(output_path, dpi=150)
    plt.close(fig)


def generate_diagrams(tmp_dir: Path) -> dict[str, Path]:
    tmp_dir.mkdir(parents=True, exist_ok=True)
    diagrams: dict[str, Path] = {}

    system_path = tmp_dir / "system_architecture.png"
    _generate_diagram(
        system_path,
        "System Architecture",
        ["NetBox (Loki) API", "Ansible Playbook\nnetbox_zabbix_sync", "Zabbix JSON-RPC API", "SMTP / Email"],
        arrows=[(0, 1), (1, 2), (1, 3)],
    )
    diagrams["system"] = system_path

    flow_path = tmp_dir / "automation_flow.png"
    _generate_diagram(
        flow_path,
        "High-Level Automation Flow",
        [
            "Playbook Start",
            "Validate Credentials",
            "Load YAML Mappings",
            "Fetch NetBox Devices",
            "Fetch Zabbix Hosts",
            "Process Devices\n(create / update)",
            "Summarize Results\n+ Email Report",
        ],
        arrows=[(i, i + 1) for i in range(6)],
    )
    diagrams["flow"] = flow_path

    mapping_path = tmp_dir / "yaml_mapping_relations.png"
    _generate_diagram(
        mapping_path,
        "YAML Configuration Relations",
        [
            "netbox_device_type_mapping.yml",
            "templates.yml",
            "template_types.yml",
            "host_groups_config.yml",
            "tags_config.yml",
        ],
        arrows=[(0, 1), (1, 2), (0, 3), (0, 4)],
    )
    diagrams["yaml_relations"] = mapping_path

    transform_path = tmp_dir / "data_transformation.png"
    _generate_diagram(
        transform_path,
        "Data Transformation",
        [
            "NetBox Device\n(dcim/devices)",
            "Python Processing\n(fetch_all_devices / processor)",
            "Zabbix Host\n(host.create / host.update)",
        ],
        arrows=[(0, 1), (1, 2)],
    )
    diagrams["transformation"] = transform_path

    return diagrams


def _add_common_intro_sections(doc) -> None:
    today = datetime.now().strftime("%Y-%m-%d")
    doc.add_heading("Zabbix-NetBox Entegrasyonu - Yonetici Dokumani", level=1)
    doc.add_paragraph(f"Tarih: {today}")

    doc.add_heading("Genel Bakis", level=2)
    doc.add_paragraph(
        "Bu dokuman, NetBox (Loki) ile Zabbix arasinda calisan otomatik envanter senkronizasyonunun "
        "yonetici duzeyinde ozetini saglar. Cozum, HMDL (Host Metadata-Driven Lifecycle) projesinin bir "
        "bileseni olarak tasarlanmistir."
    )

    doc.add_heading("Ana Bilesenler", level=2)
    components = [
        "NetBox (Loki) API - Gercek envanter ve metadata kaynagi.",
        "Ansible playbook `playbooks/netbox_zabbix_sync.yaml` ve `netbox_zabbix_sync` rolu - Tum otomasyon akisini orkestre eder.",
        "Python scriptleri - NetBox cihazlarini filtreler, eslestirir ve Zabbix icin uygun formata donusturur.",
        "Zabbix JSON-RPC API - Hedef izleme sistemi; host create/update islemleri burada gerceklesir.",
        "SMTP / Email altyapisi - Entegrasyon sonuclarini iceren ozet ve detayli raporlari ilgili ekiplere gonderir.",
    ]
    for item in components:
        doc.add_paragraph(item, style="List Bullet")


def build_docx(output_dir: Path, diagrams: dict[str, Path]) -> Path:
    doc = Document()
    _add_common_intro_sections(doc)

    doc.add_heading("Sistem Mimarisi", level=2)
    doc.add_paragraph(
        "Sistem, NetBox tarafindaki DCIM cihaz kayitlarini tek gercek kaynagi olarak kabul eder. "
        "Ansible playbook'u, NetBox API'den filtrelenmis cihaz listesini alir, Zabbix API uzerinden mevcut host "
        "bilgilerini toplar ve her bir cihaz icin create/update kararini verir."
    )
    system_img = diagrams.get("system")
    if system_img and system_img.exists():
        doc.add_picture(str(system_img), width=Inches(5.5))

    doc.add_heading("Otomasyon Akisi", level=2)
    steps = [
        "Playbook `netbox_zabbix_sync.yaml` calisir ve NetBox/Zabbix kimlik bilgilerini dogrular.",
        "Role `netbox_zabbix_sync`, YAML konfigurasyonlarini yukler ve NetBox cihazlarini filtreleyen Python scriptini calistirir.",
        "Zabbix'ten tum host listesi cekilir ve Loki_ID/hostname bazli eslestirme icin hazirlanir.",
        "Her NetBox cihazi icin Python processor, cihaz tipini, IP adresini, host group'larini ve tag'leri hesaplar.",
        "Eslestirme sonucuna gore Zabbix uzerinde yeni host olusturulur veya mevcut host guncellenir.",
        "Tum sonuc kayitlari toplanir, ozetlenir ve CSV ekli email raporu ile paylasilir.",
    ]
    for step in steps:
        doc.add_paragraph(step, style="List Number")

    flow_img = diagrams.get("flow")
    if flow_img and flow_img.exists():
        doc.add_picture(str(flow_img), width=Inches(5.5))

    doc.add_heading("YAML Konfigurasyon Dosyalari", level=2)

    doc.add_heading("netbox_device_type_mapping.yml", level=3)
    doc.add_paragraph(
        "NetBox cihaz ozelliklerini (device_role, manufacturer, model) kullanarak soyut bir DEVICE_TYPE "
        "degeri ureten kurallari tanimlar. Oncelik (priority) alanina gore ilk eslesen kural kullanilir. "
        "Bu DEVICE_TYPE degeri, secilecek Zabbix template setini ve host group mantigini dogrudan belirler."
    )

    doc.add_heading("templates.yml", level=3)
    doc.add_paragraph(
        "Her DEVICE_TYPE icin kullanilacak Zabbix template'lerini, interface tipini ve statik host group "
        "listelerini tanimlar. Ayrica API tabanli template'ler icin otomatik makro degerlerini (ornegin "
        "{HOST.IP} iceren URL'ler) konfigure eder."
    )

    doc.add_heading("template_types.yml", level=3)
    doc.add_paragraph(
        "Template tipini (snmpv2, snmpv3, agent, api) Zabbix arayuz konfigurasyonuna donusturen haritayi icerir. "
        "Bu sayede her cihaz icin dogru port, protokol ve kimlik dogrulama parametreleri otomatik olarak atanir."
    )

    doc.add_heading("host_groups_config.yml", level=3)
    doc.add_paragraph(
        "Host group degerlerinin hangi kaynaklardan ve hangi oncelik sirasiyla uretilecegini tanimlar. "
        "Mapping sonucu olusan DEVICE_TYPE, NetBox role bilgisi, lokasyon hiyerarsisi ve custom field'lar "
        "tek bir konfigurasyon dosyasi uzerinden yonetilir."
    )

    doc.add_heading("tags_config.yml", level=3)
    doc.add_paragraph(
        "Zabbix uzerinde olusacak etiketlerin (tags) NetBox alanlarindan, custom field'lardan veya "
        "hesaplanmis degerlerden nasil turetilecegini tanimlar. Ayrica NetBox `tags` dizisi, Loki_Tag_ "
        "prefix'i ile birden fazla Zabbix tag'ine genisletilebilir."
    )

    yaml_img = diagrams.get("yaml_relations")
    if yaml_img and yaml_img.exists():
        doc.add_picture(str(yaml_img), width=Inches(5.5))

    doc.add_heading("Veri Donusum Akisi", level=2)
    doc.add_paragraph(
        "NetBox'tan gelen ham cihaz nesnesi, Python processor tarafindan islenerek Zabbix host kaydina "
        "donusturulur. Bu surecte cihaz tip eslestirme, IP cozumleme, host group ve tag olusturma islemleri "
        "tamamen konfigurasyon dosyalarina dayali calisir."
    )

    transform_img = diagrams.get("transformation")
    if transform_img and transform_img.exists():
        doc.add_picture(str(transform_img), width=Inches(5.5))

    doc.add_heading("Hata Yonetimi ve Raporlama", level=2)
    doc.add_paragraph(
        "Her cihaz icin gerceklesen create/update isleminin sonucu kaydedilir. Basarisiz kayitlar icin "
        "hata mesaji ve neden bilgisi tutulur. Calisma sonunda hem ozet sayilar (eklendi, guncellendi, "
        "guncel, eklenemedi) hem de tum cihazlari iceren detayli bir CSV raporu email ile paylasilir."
    )

    output_path = output_dir / "ZABBIX_NETBOX_ADMIN_GUIDE.docx"
    doc.save(str(output_path))
    return output_path


class AdminPDF(FPDF):
    def header(self) -> None:
        self.set_font("Helvetica", "B", 14)
        self.cell(0, 10, "Zabbix-NetBox Entegrasyonu - Yonetici Dokumani", ln=1, align="C")
        self.ln(2)

    def footer(self) -> None:
        self.set_y(-15)
        self.set_font("Helvetica", "I", 8)
        page_label = f"Page {self.page_no()}"
        self.cell(0, 10, page_label, 0, 0, "C")


def _pdf_multiline(pdf: AdminPDF, text: str) -> None:
    pdf.set_font("Helvetica", size=11)
    pdf.multi_cell(0, 5, text)
    pdf.ln(2)


def build_pdf(output_dir: Path, diagrams: dict[str, Path]) -> Path:
    pdf = AdminPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    today = datetime.now().strftime("%Y-%m-%d")
    pdf.set_font("Helvetica", "", 11)
    pdf.cell(0, 8, f"Tarih: {today}", ln=1)
    pdf.ln(2)

    pdf.set_font("Helvetica", "B", 12)
    pdf.cell(0, 8, "Genel Bakis", ln=1)
    _pdf_multiline(
        pdf,
        (
            "Bu dokuman, NetBox (Loki) ile Zabbix arasinda calisan otomatik envanter senkronizasyon "
            "cozumunu yonetici bakis acisiyla ozetler. Cozum, HMDL (Host Metadata-Driven Lifecycle) "
            "projesinin bir parcasi olarak tasarlanmaktadir."
        ),
    )

    pdf.set_font("Helvetica", "B", 12)
    pdf.cell(0, 8, "Ana Bilesenler", ln=1)
    pdf.set_font("Helvetica", "", 11)
    bullets = [
        "NetBox (Loki) API - envanter ve metadata kaynagi.",
        "Ansible playbook ve rol - otomasyon akisini yurutur.",
        "Python scriptleri - cihazlarin islenmesi ve donusumu.",
        "Zabbix JSON-RPC API - izleme sisteminde host olusturma/guncelleme.",
        "SMTP / Email - ozet ve detayli raporlama.",
    ]
    for item in bullets:
        pdf.cell(5)
        pdf.cell(0, 6, f"- {item}", ln=1)
    pdf.ln(4)

    if "system" in diagrams and diagrams["system"].exists():
        pdf.set_font("Helvetica", "B", 12)
        pdf.cell(0, 8, "Sistem Mimarisi", ln=1)
        pdf.image(str(diagrams["system"]), w=170)
        pdf.ln(4)

    pdf.add_page()
    pdf.set_font("Helvetica", "B", 12)
    pdf.cell(0, 8, "Otomasyon Akisi", ln=1)
    steps = [
        "Playbook baslatilir ve kimlik bilgileri dogrulanir.",
        "YAML konfigurasyonlari yuklenir ve NetBox cihazlari filtrelenir.",
        "Zabbix'ten mevcut host listesi cekilir.",
        "Her cihaz icin tip, IP, host gruplari ve tag'ler hesaplanir.",
        "Zabbix'te yeni host olusturulur veya mevcut host guncellenir.",
        "Sonuclar ozetlenir ve email/CSV raporu gonderilir.",
    ]
    pdf.set_font("Helvetica", "", 11)
    for idx, step in enumerate(steps, start=1):
        pdf.cell(5)
        pdf.cell(0, 6, f"{idx}. {step}", ln=1)
    pdf.ln(4)
    if "flow" in diagrams and diagrams["flow"].exists():
        pdf.image(str(diagrams["flow"]), w=170)
        pdf.ln(4)

    pdf.add_page()
    pdf.set_font("Helvetica", "B", 12)
    pdf.cell(0, 8, "YAML Konfigurasyon Dosyalari", ln=1)
    pdf.ln(2)

    sections = [
        (
            "netbox_device_type_mapping.yml",
            "NetBox cihaz ozelliklerini kullanarak soyut DEVICE_TYPE degerini ureten kurallari tanimlar. "
            "Oncelik sirasi ile ilk eslesme kullanilir ve bu sonuc sonraki tum eslestirme zincirini tetikler.",
        ),
        (
            "templates.yml",
            "Her DEVICE_TYPE icin uygulanacak Zabbix template setini, interface tipini ve ek host group "
            "bilgilerini tanimlar. Ayrica bazi API tabanli template'ler icin makro degerleri ayarlanir.",
        ),
        (
            "template_types.yml",
            "Template tipini (snmpv2, snmpv3, agent, api) Zabbix arayuz konfigurasyonuna ceviren sozluk gorevi gorur.",
        ),
        (
            "host_groups_config.yml",
            "Host group degerlerinin hangi NetBox alanlarindan ve hangi oncelikle uretilecegini belirler.",
        ),
        (
            "tags_config.yml",
            "Zabbix tag'lerinin NetBox alanlari, custom field'lar ve NetBox `tags` dizisinden nasil turetilecegini tanimlar.",
        ),
    ]
    for title, body in sections:
        pdf.set_font("Helvetica", "B", 11)
        pdf.cell(0, 7, title, ln=1)
        _pdf_multiline(pdf, body)

    if "yaml_relations" in diagrams and diagrams["yaml_relations"].exists():
        pdf.image(str(diagrams["yaml_relations"]), w=170)
        pdf.ln(4)

    pdf.add_page()
    pdf.set_font("Helvetica", "B", 12)
    pdf.cell(0, 8, "Veri Donusum Akisi", ln=1)
    _pdf_multiline(
        pdf,
        (
            "NetBox cihaz nesnesi, Python tabanli isleme katmani tarafindan DEVICE_TYPE, IP, host group ve "
            "tag'lere donusturulur. Bu sonuc, Zabbix API uzerinden host.create veya host.update cagrisina "
            "girdi saglar. Tum mantik, YAML konfigurasyon dosyalari ile kontrol edilmektedir."
        ),
    )
    if "transformation" in diagrams and diagrams["transformation"].exists():
        pdf.image(str(diagrams["transformation"]), w=170)
        pdf.ln(4)

    pdf.set_font("Helvetica", "B", 12)
    pdf.cell(0, 8, "Hata Yonetimi ve Raporlama", ln=1)
    _pdf_multiline(
        pdf,
        (
            "Islem sonunda her cihaz icin durum (eklendi, guncellendi, guncel, eklenemedi) bilgisi saklanir. "
            "Basarisiz kayitlarin ayrintili nedenleri e-posta iceriginde ve CSV raporda yer alir; boylece operasyon "
            "ekipleri sorunlu kayitlara hizla mudahale edebilir."
        ),
    )

    output_path = output_dir / "ZABBIX_NETBOX_ADMIN_GUIDE.pdf"
    pdf.output(str(output_path))
    return output_path


def main() -> None:
    output_dir = _ensure_output_dir()
    tmp_dir = output_dir / "tmp_admin_doc_diagrams"

    diagrams = generate_diagrams(tmp_dir)
    docx_path = build_docx(output_dir, diagrams)
    pdf_path = build_pdf(output_dir, diagrams)

    print(f"DOCX generated at: {docx_path}")
    print(f"PDF generated at: {pdf_path}")


if __name__ == "__main__":
    main()

